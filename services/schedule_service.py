import asyncio
import os
import re
from datetime import date, timedelta, datetime
from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass, field

import aiofiles
import aiohttp
from docx import Document
from docx.table import _Cell
from loguru import logger
from bs4 import BeautifulSoup

from core.constants import Weekdays, IMOMI_SCHEDULES_PAGE_URL, IMOMI_ROOT_URL
from core.utils import is_bottom_week, singleton


@singleton
class ScheduleService:
    def __init__(self):
        logger.info(f"Create instance of {type(self).__name__}")
        self.update_delay = 28800  # every 8 hours runs update_schedule_task
        self.location = Path("Туризм.docx")
        self.update_schedule_task = asyncio.create_task(self._update_schedule_task())
        self.parser: Optional[DailyScheduleParser] = None
        logger.info(f"{type(self).__name__} created successfully")

    async def fetch_schedule_file(self):
        async with aiohttp.ClientSession() as session:
            link = await ScheduleDocDownloader.get_link(session)
            logger.debug(f"Fetching {link}")
            await ScheduleDocDownloader.fetch(session, link, self.location, self.location.name)
            logger.debug(
                "Schedule downloaded to {loc}. Next update will be at {next_launch_time} ({update_delay}s)".format(
                    loc=self.location,
                    next_launch_time=datetime.now() + timedelta(seconds=self.update_delay),
                    update_delay=self.update_delay,
                )
            )
        if not self.parser:
            self.parser = DailyScheduleParser(self.location)

    async def _update_schedule_task(self):
        while True:
            logger.debug("Starting update schedule task")
            await self.fetch_schedule_file()
            await asyncio.sleep(self.update_delay)

    async def tomorrow_schedule(self, grade: int) -> dict[str, Optional[str]]:
        return await self.daily_schedule(date.today() + timedelta(days=1), grade)

    async def today_schedule(self, grade: int) -> dict[str, Optional[str]]:
        return await self.daily_schedule(date.today(), grade)

    async def daily_schedule(self, d: date, grade: int) -> dict[str, Optional[str]]:
        logger.debug(f"Call daily_schedule(d={d}, grade={grade})")
        mapping = await self.parser.lesson_slots_mapping(grade, Weekdays.from_date(d))
        schedule = {}
        if is_bottom_week(d):
            schedule = {t: ls.bottom_week_lesson for t, ls in mapping.items()}
        else:
            schedule = {t: ls.top_week_lesson for t, ls in mapping.items()}
        logger.debug(f"daily_schedule(d={d}, grade={grade}): {schedule}")
        return schedule


@dataclass
class LessonSlot:
    top_week_lesson: Optional[str] = None
    bottom_week_lesson: Optional[str] = None


class DailyScheduleParser:
    def __init__(self, filepath: Path):
        logger.info(f"Init {type(self).__name__}")
        self.filepath = filepath
        self.docx = Document(self.filepath)
        logger.info(f"{type(self).__name__} initialized successfully")

    async def cell_pairs(self, grade: int, day: Weekdays) -> list[Tuple[_Cell]]:
        grade_to_table_map = {grade + 1: self.docx.tables[grade] for grade, _ in enumerate(self.docx.tables)}
        allday_cells: list[Tuple[_Cell]] = {
            grade: [row.cells for row in table.rows] for grade, table in grade_to_table_map.items()
        }[grade]

        cells = []
        adding_cells = False
        for cell in allday_cells:
            maybe_day = cell[0].text
            if maybe_day == day.value:
                adding_cells = True
            elif maybe_day in [d.value for d in Weekdays] and maybe_day != day.value:
                adding_cells = False

            if adding_cells:
                cells.append(cell)

        return cells[1:]

    async def lesson_slots_mapping(self, grade: int, day_of_week: Weekdays) -> dict[str, LessonSlot]:
        if day_of_week == Weekdays.Sunday:
            return {}

        schedule: dict[str, LessonSlot] = {}
        cells = await self.cell_pairs(grade, day_of_week)

        for time_cell, lesson_info_cell in cells:
            time: str = time_cell.text
            lesson_info = lesson_info_cell.text if lesson_info_cell.text.replace(" ", "") != "" else None

            if not schedule.get(time):
                schedule[time] = LessonSlot(top_week_lesson=lesson_info, bottom_week_lesson=lesson_info)
            else:
                schedule[time].bottom_week_lesson = lesson_info

        return schedule


class ScheduleDocDownloader:
    @staticmethod
    async def get_link(session: aiohttp.ClientSession) -> str:
        async with session.get(IMOMI_SCHEDULES_PAGE_URL, timeout=0) as response:
            bs = BeautifulSoup(await response.text(), features="html.parser")
            panels = bs.find_all("div", {"class": "panel-body"})
            bachelor_panel = next((p for p in panels if "Бакалавриат" in p.text), None)
            found_link = bachelor_panel.find_next("a", text=re.compile(".*Туризм.*")).attrs.get("href")
            return IMOMI_ROOT_URL + found_link

    @staticmethod
    async def fetch(session: aiohttp.ClientSession, link: str, location: Path, name: str) -> None:
        """
        https://stackoverflow.com/questions/58804285/asynchronous-download-of-files
        """
        async with session.get(link, timeout=0) as response:
            location = location.absolute()
            logger.debug(f"Fetching {name}")
            try:
                file = await aiofiles.open(location, mode="wb")
                await file.write(await response.content.read())
                await file.close()
                logger.debug(f"{name} downloaded")
            except FileNotFoundError:
                loc = "".join((r"/".join((str(location).split("/")[:-1])), "/"))
                command = " ".join(("mkdir -p", loc))
                os.system(command)
                file = await aiofiles.open(location, mode="wb")
                await file.write(await response.content.read())
                await file.close()
                logger.debug(f"{name} downloaded")
