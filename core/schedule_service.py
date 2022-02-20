from datetime import date, timedelta
from pathlib import Path
from typing import Optional, Dict, List, Tuple
from dataclasses import dataclass
from docx import Document
from docx.table import _Cell
from loguru import logger

from core.constants import Weekdays
from core.utils import is_bottom_week


class ScheduleService:
    def __init__(self):
        logger.info("Init ScheduleService")
        self.schedule_parser = DailyScheduleParser(Path("Туризм.docx"))
        logger.info("ScheduleService initialized successfully")

    def tomorrow_schedule(self, grade: int) -> Dict[str, Optional[str]]:
        return self.daily_schedule(date.today() + timedelta(days=1), grade)

    def today_schedule(self, grade: int) -> Dict[str, Optional[str]]:
        return self.daily_schedule(date.today(), grade)

    def daily_schedule(self, d: date, grade: int) -> Dict[str, Optional[str]]:
        logger.debug(f"Call daily_schedule(d={d}, grade={grade})")
        mapping = self.schedule_parser.lesson_slots_mapping(grade, Weekdays.from_date(d))
        schedule = {}
        if is_bottom_week(d):
            schedule = {t: ls.bottom_week_lesson for t, ls in mapping.items()}
        else:
            schedule = {t: ls.top_week_lesson for t, ls in mapping.items()}
        logger.debug(f"daily_schedule(d={d}, grade={grade}): {schedule}")
        return schedule


@dataclass
class LessonSlot:
    top_week_lesson: Optional[str] = None
    bottom_week_lesson: Optional[str] = None


class DailyScheduleParser:
    def __init__(self, filepath: Path):
        logger.debug(f"Init DailyScheduleParser(file={filepath.name})")
        self.filepath = filepath
        self._docx = Document(self.filepath)
        logger.debug(f"DailyScheduleParser(file={filepath.name}) initialized successfully")

    def cell_pairs(self, grade: int, day: Weekdays) -> List[Tuple[_Cell]]:
        grade_to_table_map = {grade + 1: self._docx.tables[grade] for grade, _ in enumerate(self._docx.tables)}
        allday_cells: List[Tuple[_Cell]] = {
            grade: [row.cells for row in table.rows] for grade, table in grade_to_table_map.items()
        }[grade]

        cells = []
        adding_cells = False
        for cell in allday_cells:
            maybe_day = cell[0].text
            if maybe_day == day.value:
                adding_cells = True
            elif maybe_day in [d.value for d in Weekdays] and maybe_day != day.value:
                adding_cells = False

            if adding_cells:
                cells.append(cell)

        return cells[1:]

    def lesson_slots_mapping(self, grade: int, day_of_week: Weekdays) -> Dict[str, LessonSlot]:
        if day_of_week == Weekdays.Sunday:
            return {}

        schedule: Dict[str, LessonSlot] = {}
        cells = self.cell_pairs(grade, day_of_week)

        for time_cell, lesson_info_cell in cells:
            time: str = time_cell.text
            lesson_info = lesson_info_cell.text if lesson_info_cell.text.replace(" ", "") != "" else None

            if not schedule.get(time):
                schedule[time] = LessonSlot(top_week_lesson=lesson_info, bottom_week_lesson=lesson_info)
            else:
                schedule[time].bottom_week_lesson = lesson_info

        return schedule
