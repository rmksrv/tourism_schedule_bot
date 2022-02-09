from pathlib import Path
from typing import Optional, Dict, List, Tuple
from dataclasses import dataclass
from docx import Document
from docx.table import _Cell

from core.constants import Days


@dataclass
class LessonSlot:
    top_week_lesson: Optional[str] = None
    bottom_week_lesson: Optional[str] = None


class DaySchedule:
    def __init__(self, grade: int, day: Days):
        self.grade = grade
        self.day = day

    def formatted_schedule(self, filepath: Path) -> Dict[str, LessonSlot]:
        schedule: Dict[str, LessonSlot] = {}
        cells = self.cells_from_docx(filepath)

        for time_cell, lesson_info_cell in cells:
            time: str = time_cell.text
            lesson_info = lesson_info_cell.text if lesson_info_cell.text != "" else None

            if not schedule.get(time):
                schedule[time] = LessonSlot(
                    top_week_lesson=lesson_info, bottom_week_lesson=lesson_info
                )
            else:
                schedule[time].bottom_week_lesson = lesson_info

        return schedule

    def cells_from_docx(self, filepath: Path) -> List[Tuple[_Cell]]:
        docx = Document(filepath)
        grade_to_table_map = {
            grade + 1: docx.tables[grade] for grade, _ in enumerate(docx.tables)
        }
        allday_cells: List[Tuple[_Cell]] = {
            grade: [row.cells for row in table.rows]
            for grade, table in grade_to_table_map.items()
        }[self.grade]

        cells = []
        adding_cells = False
        for cell in allday_cells:
            maybe_day = cell[0].text
            if maybe_day == self.day.value:
                adding_cells = True
            elif maybe_day in [d.value for d in Days] and maybe_day != self.day.value:
                adding_cells = False

            if adding_cells:
                cells.append(cell)

        return cells[1:]
